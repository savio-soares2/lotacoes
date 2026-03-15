from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm


COLUMN_LABELS = {
    "identificador": "Identificador",
    "nome": "Nome",
    "unidade": "Unidade",
    "opcao_contemplada": "Opção Contemplada",
    "criterio_aplicado": "Critério Aplicado",
    "tempo_servico_dias": "Tempo de Serviço (dias)",
    "idade_dias": "Idade (dias)",
    "data_admissao": "Data de Admissão",
    "data_nascimento": "Data de Nascimento",
    "opcao_em_analise": "Opção em Análise",
    "vagas_em_disputa": "Vagas em Disputa",
    "motivo": "Motivo",
    "opcao_1": "1ª Opção",
    "opcao_2": "2ª Opção",
    "opcao_3": "3ª Opção",
    "vagas_totais": "Vagas Totais",
    "lotados_automatico": "Lotados Automaticamente",
    "reservadas_para_desempate_manual": "Reservadas para Desempate Manual",
    "vagas_restantes": "Vagas Restantes",
    "unidade_solicitada": "Unidade Solicitada",
    "opcao_solicitada": "Opção Solicitada",
    "resultado_na_unidade": "Resultado na Unidade Solicitada",
    "unidade_lotada_final": "Unidade de Lotação Final",
    "opcao_contemplada_final": "Opção Contemplada Final",
    "criterio_resultado_final": "Critério do Resultado Final",
    "detalhamento_resultado": "Detalhamento do Resultado",
}

SUMMARY_LABELS = {
    "total_servidores": "Total de Servidores",
    "lotados_automaticamente": "Lotados Automaticamente",
    "em_desempate_manual": "Em Desempate Manual",
    "nao_lotados": "Não Lotados",
    "percentual_desempate_manual": "Percentual em Desempate Manual",
    "data_referencia_criterios": "Data de Referência dos Critérios",
}


def _add_table_from_rows(document: Document, title: str, rows: list[dict[str, Any]]) -> None:
    document.add_heading(title, level=2)
    if not rows:
        document.add_paragraph("Sem registros.")
        return

    headers = [h for h in rows[0].keys() if h != "identificador"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = str(COLUMN_LABELS.get(h, h))

    for row in rows:
        cells = table.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = str(row.get(h, ""))



def build_report(result: dict[str, Any]) -> bytes:
    doc = Document()
    
    # Configurar orientação Paisagem (Landscape) e margens estreitas
    section = doc.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    doc.add_heading("Relatório de Lotação de Servidores", level=1)
    doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

    doc.add_heading("Critérios Utilizados", level=2)
    doc.add_paragraph("1. Tempo de serviço (data de admissão mais antiga).")
    doc.add_paragraph("2. Idade (data de nascimento mais antiga).")
    doc.add_paragraph("3. Distância residência-unidade, apenas para os casos de desempate manual.")

    resumo = result.get("resumo", {})
    doc.add_heading("Resumo Geral", level=2)
    for k, v in resumo.items():
        label = SUMMARY_LABELS.get(k, k)
        if k == "percentual_desempate_manual":
            doc.add_paragraph(f"{label}: {v}%")
        else:
            doc.add_paragraph(f"{label}: {v}")

    _add_table_from_rows(doc, "Resumo por Unidade", result.get("resumo_unidades", []))

    doc.add_heading("Solicitações por Unidade", level=2)
    solicitacoes_por_unidade = result.get("solicitacoes_por_unidade", {})
    if not solicitacoes_por_unidade:
        doc.add_paragraph("Sem registros de solicitações por unidade.")
    else:
        for unidade, rows in solicitacoes_por_unidade.items():
            doc.add_heading(f"Unidade: {unidade}", level=3)
            if not rows:
                doc.add_paragraph("Sem solicitações para esta unidade.")
                continue

            doc.add_paragraph(
                "A tabela abaixo inclui todos que pediram esta unidade, inclusive os não lotados nela, com motivo detalhado."
            )
            headers = [h for h in rows[0].keys() if h != "identificador"]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"

            header_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                header_cells[i].text = str(COLUMN_LABELS.get(h, h))

            for row in rows:
                cells = table.add_row().cells
                for i, h in enumerate(headers):
                    cells[i].text = str(row.get(h, ""))

    _add_table_from_rows(doc, "Casos para Desempate Manual", result.get("desempate_manual", []))
    _add_table_from_rows(doc, "Não Lotados", result.get("nao_lotados", []))

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
